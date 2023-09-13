import os

import openpyxl
# from tkinter import filedialog

# import matplotlib

# matplotlib.use("agg")
# import matplotlib.pyplot as plt
# import docx
from docx import Document
from openpyxl import load_workbook

print("nihao")

# def change(mainForm, control, text):
#     mainForm[control].text = text


text1 = " "
text2 = " "
text3 = " "
text4 = " "
text5 = " "
url = "C:/Users/ChuanZhou/Desktop/ces"
mode = 0


def change(
        mainForm):  # liuzhe yong zuo chuang ru chang shu de ru kou chuang ru hou zai jie mian cheng xu li mian diao yong change
    # suo yi shi ji shang zhe ge hangshu jiu shi gao zhi zai jiemian li mian diao yong zhi hou yao zuo xie shen me
    # mainForm["edit"].text = text
    global text1
    global text2
    global text3
    global text4
    global text5
    global url  # biao shi shi yong de shi quan ju bian liang
    text1 = mainForm["edit"].text
    text2 = mainForm["edit2"].text
    text3 = mainForm["edit3"].text
    text4 = mainForm["edit4"].text
    text5 = mainForm["edit5"].text
    url = mainForm["edit6"].text
    # print(text1)


def xieru(mainForm):
    hangshu = 1
    global url
    global mode
    path = url
    files_list = os.listdir(path)
    baochun_path = url + '/' + 'turnfile_rain'
    if os.path.exists(baochun_path):
        # 前面已经有了输出文件夹所以嘛有要的excel了啦所以下方打开地址
        wb = load_workbook(baochun_path + '/' + 'num.xlsx')#xie ru excel de tou nei rong
        sheet = wb.active
        sheet.cell(row=1, column=1).value = '学号'
        sheet.cell(row=1, column=2).value = '姓名'
        sheet.cell(row=1, column=3).value = '分数'
        wb.save(baochun_path + '/' + 'num.xlsx')
    else:
        os.mkdir(baochun_path)
        #前面刚建立输出文件夹所以嘛肯定没有要的excel了啦所以下方新建
        wb = openpyxl.Workbook() #xing jian excel
        sheet = wb.active
        sheet.cell(row=1, column=1).value = '学号'
        sheet.cell(row=1, column=2).value = '姓名'
        sheet.cell(row=1, column=3).value = '分数'
        wb.save(baochun_path + '/' + 'num.xlsx')

    for file_name in files_list:
        # 判断文件类型是否在文件名中
        if '.docx' in file_name:
            document = Document(url + '/' + file_name)
            print(url)
            # paragraphs = document.paragraphs

            table = document.tables
            ro = table[0].rows[5]  # chao zuo di yi ge biao ge de di liu hang nei rong
            chengji_str = ro.cells[1].text  # huo qu di er ge dang yuan ge de nei rong
            chengji = int(chengji_str, 10)
            name = table[0].rows[2].cells[1].text
            xuehao = table[0].rows[3].cells[1].text
            # print(name)
            # print(xuehao)

            print(chengji_str)

            for table in document.tables:  # pang duan du yu bu tong cheng ji xie bu tong ping yu   以及遍历表格找到指定位置进行写入
                for ro in table.rows:
                    for cell in ro.cells:
                        # print(cell.text)
                        # if(cell.text == )
                        if "指导老师评语：" in cell.text:
                            print("找到评语位置准备写入")
                            if chengji >= 90:
                                print("da yu 90")
                                cell.text = cell.text.replace('指导老师评语：', '指导老师评语：\n' + text1)
                                mode = 1
                            elif chengji >= 80:
                                print("da yu 80 ")  # pang duan cheng ji suo chu de duan wei
                                cell.text = cell.text.replace('指导老师评语：', '指导老师评语：\n' + text2)
                                mode = 2
                            elif chengji >= 70:
                                print("da yu 70 ")  # pang duan cheng ji suo chu de duan wei
                                cell.text = cell.text.replace('指导老师评语：', '指导老师评语：\n' + text3)
                                mode = 3
                            elif chengji >= 60:
                                print("da yu 60")  # pang duan cheng ji suo chu de duan wei
                                cell.text = cell.text.replace('指导老师评语：', '指导老师评语：\n' + text4)
                                mode = 4
                            else:
                                print("这个逼不及格")
                                cell.text = cell.text.replace('指导老师评语：', '指导老师评语：\n' + text5)
                                mode = 5
                            # cell.text = cell.text.replace('指导老师评语：', '指导老师评语：' + text1)
            document.save(baochun_path + '/' + file_name)
            print(file_name + "word写入完成")
            hangshu = hangshu+1
            print(file_name + "提取对应内容到excel")
            wb = load_workbook(baochun_path+'/'+'num.xlsx')
            sheet = wb.active
            sheet.cell(row=hangshu, column=1).value = xuehao
            sheet.cell(row=hangshu, column=2).value = name
            sheet.cell(row=hangshu, column=3).value = chengji
            wb.save(baochun_path+'/'+'num.xlsx')
            print("excel写入完成")

    # for para in document.paragraphs:
    #     if '计算机通信网络课程设计报告' in para.text:
    #         para.text = para.text.replace('计算机通信网络课程设计报告', text1)
    # document.save("2020_end.docx")
    # table = document.tables
    # ro = table[0].rows[5]  # chao zuo di yi ge biao ge de di liu hang nei rong
    # chengji_str = ro.cells[1].text  # huo qu di er ge dang yuan ge de nei rong
    # chengji = int(chengji_str, 10)
    # print(chengji_str)
    # if chengji >= 90:
    #     print("da yu 90")
    # elif chengji >= 80:
    #     print("da yu 80 xiao yu 90")
    # table = document.tables
    # print(table[1].rows[0].cells[0].text)
    print("成功运行end")


# def get_path(mainForm):
#     global url
#     # url = filedialog.askdirectory()
#     print(url)


# get_path(1)
# xieru(1)
# doc = docx.Document(doc_name)
# document = Document(url)
# paragraphs = document.paragraphs

# for para in document.paragraphs:
#   print(para.text)

# for para in document.paragraphs:
#    if '计算机通信网络课程设计报告' in para.text:
#        para.text = para.text.replace('计算机通信网络课程设计报告', '计算机通信网络课程设计报告35241564165')
# paragraphs[0].add_run("我在这里我在这里呢快来我在这里")
# print(url)

# document = Document(url)
# print(url)
# paragraphs = document.paragraphs
# for table in document.tables:
#     for ro in table.rows:
#         for cell in ro.cells:
#             print(cell.text)
#             if "指导老师评语：" in cell.text:
#                 print("youde")
#                 cell.text = cell.text.replace('指导老师评语：', '指导老师评语：符合送i发hi文化')
# document.save("2020_end.docx")
